import sys
import re
import os
import shutil
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QFileDialog,
    QVBoxLayout, QMessageBox
)


class FileValidator:
    def __init__(self, doc, filename):
        self.doc = doc
        self.filename = filename
        self.errors = []
        self.code_pattern = re.compile(r'^[A-ZА-Я]+\s*-\s*\d+$')
        self.range_pattern = re.compile(r'^\d+-\d+$')

    def validate(self):
        self.validate_competency_codes()
        self.validate_task_numbers()
        return not self.errors

    def validate_competency_codes(self):
        if not self.doc.tables:
            self.errors.append({"Тип ошибки": "Нет таблиц в документе", "Строка": ""})
            return
        first_table = self.doc.tables[0]
        for row in first_table.rows[1:]:
            code = row.cells[0].text.strip()
            if not self.code_pattern.match(code):
                self.errors.append({
                    "Тип ошибки": "Неверный формат кода компетенции",
                    "Строка": code
                })

    def validate_task_numbers(self):
        first_table = self.doc.tables[0]
        for row in first_table.rows[1:]:
            num_text = row.cells[-1].text.strip()
            if not self.range_pattern.match(num_text):
                self.errors.append({
                    "Тип ошибки": "Неверный формат диапазона номеров заданий",
                    "Строка": num_text
                })


class CompetencyExtractor(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Разделение компетенций Word-файлов")
        self.setGeometry(300, 300, 400, 150)

        self.button = QPushButton("Выбрать файлы Word (.docx)", self)
        self.button.clicked.connect(self.process_files)

        layout = QVBoxLayout()
        layout.addWidget(self.button)
        self.setLayout(layout)

    def process_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Выберите Word-файлы", "", "Word Files (*.docx)")
        if not files:
            return

        result_dir = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения результатов")
        if not result_dir:
            QMessageBox.warning(self, "Отмена", "Операция отменена.")
            return

        success_dir = os.path.join(result_dir, "Успешно разрезанные ФОС")
        failed_dir = os.path.join(result_dir, "Не форматные исходные файлы ФОС")
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(failed_dir, exist_ok=True)

        wb = Workbook()
        ws = wb.active
        ws.append(["Наименование файла", "Тип ошибки", "Строка с ошибкой"])

        for file_path in files:
            filename = os.path.basename(file_path)
            doc = Document(file_path)

            for table in doc.tables:
                self.set_table_borders(table)

            validator = FileValidator(doc, filename)
            if validator.validate():
                try:
                    self.process_file(doc, filename, success_dir)
                except Exception as e:
                    ws.append([filename, "Ошибка обработки компетенций", str(e)])
                    shutil.copy(file_path, failed_dir)
            else:
                for err in validator.errors:
                    ws.append([filename, err["Тип ошибки"], err["Строка"]])
                shutil.copy(file_path, failed_dir)

        wb.save(os.path.join(failed_dir, "Отчет_ошибок.xlsx"))
        QMessageBox.information(self, "Готово", "Обработка завершена.")

    def process_file(self, original_doc, original_filename, output_dir):
        tables = original_doc.tables
        first_table = tables[0]
        second_table = tables[1]

        # Собираем все компетенции
        competencies = []
        for row in first_table.rows[1:]:
            code = row.cells[0].text.strip().replace(" ", "")
            competencies.append({'code': code, 'row': row})

        # Находим начало Перечня заданий
        per_list_start = None
        for i, para in enumerate(original_doc.paragraphs):
            if "Перечень заданий" in para.text:
                per_list_start = i
                break

        if per_list_start is None:
            raise Exception("Раздел 'Перечень заданий' не найден")

        # Берём все элементы после Перечня заданий
        document_elements = list(original_doc.element.body[per_list_start + 1:])
        competency_code_pattern = re.compile(r'^[A-ZА-Я]+\s*-\s*\d+', re.IGNORECASE)

        # Составляем список номеров заданий из второй таблицы
        task_rows = []
        for row in second_table.rows[1:]:
            num_text = row.cells[0].text.strip()
            try:
                num = int(num_text.replace(".", "").strip())
                task_rows.append((num, row))
            except ValueError:
                continue

        # Проходим по компетенциям
        for comp in competencies:
            # Выделяем блок элементов для компетенции
            copying = False
            current_elements = []

            for el in document_elements:
                if el.tag.endswith('p'):
                    p = self._element_to_paragraph(el)
                    txt = p.text.strip().replace(" ", "")
                    if txt == comp['code']:
                        copying = True
                        continue
                    elif competency_code_pattern.match(txt) and copying:
                        break
                if copying:
                    current_elements.append(el)

            # Проверяем количество инструкций
            instruction_pattern = re.compile(r'^(\d+)\.\s*Инструкция:')
            instruction_numbers = []
            for el in current_elements:
                if el.tag.endswith('p'):
                    p = self._element_to_paragraph(el)
                    text = p.text.strip()
                    m = instruction_pattern.match(text)
                    if m:
                        instruction_numbers.append(int(m.group(1)))

            num_text = comp['row'].cells[-1].text.strip()
            m = re.match(r'(\d+)-(\d+)', num_text)
            if not m:
                raise Exception(
                    f"Компетенция {comp['code']}: неверный формат диапазона номеров '{num_text}' (должно быть например - 1-16)"
                )
            start = int(m.group(1))
            end = int(m.group(2))
            expected_count = end - start + 1

            if len(instruction_numbers) != expected_count:
                raise Exception(
                    f"Компетенция {comp['code']}: количество заданий ({len(instruction_numbers)}) "
                    f"не совпадает с ожидаемым ({expected_count})"
                )

            # Создаем новый документ для компетенции
            new_doc = Document()

            # Первая таблица
            t1 = new_doc.add_table(rows=1, cols=len(first_table.columns))
            for i, c in enumerate(first_table.rows[0].cells):
                t1.rows[0].cells[i].text = c.text
            r = t1.add_row()
            for i, c in enumerate(comp['row'].cells):
                r.cells[i].text = c.text
            self.set_table_borders(t1)

            new_doc.add_paragraph("\n")

            # Вторая таблица
            nums = list(range(start, end + 1))
            t2 = new_doc.add_table(rows=1, cols=len(second_table.columns))
            for i, c in enumerate(second_table.rows[0].cells):
                t2.rows[0].cells[i].text = c.text
            for n, r in task_rows:
                if n in nums:
                    row = t2.add_row()
                    for i, c in enumerate(r.cells):
                        row.cells[i].text = c.text
            self.set_table_borders(t2)

            new_doc.add_paragraph("\n")
            heading = new_doc.add_paragraph("Перечень заданий")
            heading.style = 'Heading 2'
            new_doc.add_paragraph("\n")

            # Добавляем элементы блока
            for el in current_elements:
                new_doc.element.body.append(deepcopy(el))

            filename = f"{comp['code']}_{original_filename}"
            new_doc.save(os.path.join(output_dir, filename))

    def set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        tblBorders = OxmlElement('w:tblBorders')

        borders = [
            ('top', 'single', 4, '000000'),
            ('left', 'single', 4, '000000'),
            ('bottom', 'single', 4, '000000'),
            ('right', 'single', 4, '000000'),
            ('insideH', 'single', 4, '000000'),
            ('insideV', 'single', 4, '000000')
        ]

        for border_type, border_style, border_size, border_color in borders:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _element_to_paragraph(self, element):
        from docx.oxml import parse_xml
        from docx.text.paragraph import Paragraph
        return Paragraph(parse_xml(element.xml), None)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = CompetencyExtractor()
    window.show()
    sys.exit(app.exec_())
