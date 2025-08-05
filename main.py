import sys
import re
import os
import shutil
import io
from collections import defaultdict
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm, Inches, Length, Mm, Emu
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import Workbook
from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QFileDialog, QVBoxLayout,
    QMessageBox, QHBoxLayout, QProgressBar, QLabel, QLineEdit, QTabWidget, QFormLayout
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont


class FileValidator:
    def __init__(self, doc, filename):
        self.doc = doc
        self.filename = filename
        self.errors = []
        self.code_pattern = re.compile(r'^[A-ZА-Я]+\s*-\s*\d+$')
        self.range_pattern = re.compile(r'^\d+-\d+$')

    def validate(self):
        self.validate_competency_codes()
        self.validate_task_numbers()
        return not self.errors

    def validate_competency_codes(self):
        if not self.doc.tables:
            self.errors.append({"Тип ошибки": "Нет таблиц в документе", "Строка": ""})
            return
        first_table = self.doc.tables[0]
        for row in first_table.rows[1:]:
            code = row.cells[0].text.strip()
            if not self.code_pattern.match(code):
                self.errors.append({
                    "Тип ошибки": "Неверный формат кода компетенции",
                    "Строка": code
                })

    def validate_task_numbers(self):
        first_table = self.doc.tables[0]
        for row in first_table.rows[1:]:
            num_text = row.cells[-1].text.strip()
            if not self.range_pattern.match(num_text):
                self.errors.append({
                    "Тип ошибки": "Неверный формат диапазона номеров заданий",
                    "Строка": num_text
                })


class CompetencySplitterTab(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.button = QPushButton("Выбрать файлы Word (.docx)", self)
        self.button.clicked.connect(self.process_files)

        self.progress = QProgressBar(self)
        self.progress.setAlignment(Qt.AlignCenter)
        self.progress.setVisible(False)

        self.status_label = QLabel("", self)
        self.status_label.setAlignment(Qt.AlignCenter)

        layout.addWidget(self.button)
        layout.addWidget(self.progress)
        layout.addWidget(self.status_label)
        self.setLayout(layout)

    def process_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Выберите Word-файлы", "", "Word Files (*.docx)")
        if not files:
            return

        result_dir = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения результатов")
        if not result_dir:
            QMessageBox.warning(self, "Отмена", "Операция отменена.")
            return

        success_dir = os.path.join(result_dir, "Успешно разрезанные ФОС")
        failed_dir = os.path.join(result_dir, "Не форматные исходные файлы ФОС")
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(failed_dir, exist_ok=True)

        wb = Workbook()
        ws = wb.active
        ws.append(["Наименование файла", "Тип ошибки", "Строка с ошибкой"])

        self.progress.setMaximum(len(files))
        self.progress.setValue(0)
        self.progress.setVisible(True)
        self.status_label.setText("Начата обработка файлов...")

        QApplication.processEvents()

        for idx, file_path in enumerate(files, start=1):
            filename = os.path.basename(file_path)
            self.status_label.setText(f"Обрабатывается файл: {filename}")
            QApplication.processEvents()

            doc = Document(file_path)

            for table in doc.tables:
                self.set_table_borders(table)

            validator = FileValidator(doc, filename)
            if validator.validate():
                try:
                    self.process_file(doc, filename, success_dir)
                except Exception as e:
                    ws.append([filename, "Ошибка обработки компетенций", str(e)])
                    shutil.copy(file_path, failed_dir)
            else:
                for err in validator.errors:
                    ws.append([filename, err["Тип ошибки"], err["Строка"]])
                shutil.copy(file_path, failed_dir)

            self.progress.setValue(idx)
            QApplication.processEvents()

        wb.save(os.path.join(failed_dir, "Отчет_ошибок.xlsx"))
        self.progress.setVisible(False)
        self.status_label.setText("Обработка завершена.")
        QMessageBox.information(self, "Готово", "Все файлы обработаны.")

    def process_file(self, original_doc, original_filename, output_dir):
        tables = original_doc.tables
        first_table = tables[0]
        second_table = tables[1]

        competencies = []
        for row in first_table.rows[1:]:
            code = row.cells[0].text.strip().replace(" ", "")
            competencies.append({'code': code, 'row': row})

        per_list_start = None
        for i, para in enumerate(original_doc.paragraphs):
            if "Перечень заданий" in para.text:
                per_list_start = i
                break

        if per_list_start is None:
            raise Exception("Раздел 'Перечень заданий' не найден")

        document_elements = list(original_doc.element.body[per_list_start + 1:])
        competency_code_pattern = re.compile(r'^[A-ZА-Я]+\s*-\s*\d+', re.IGNORECASE)

        task_rows = []
        for row in second_table.rows[1:]:
            num_text = row.cells[0].text.strip()
            try:
                num = int(num_text.replace(".", "").strip())
                task_rows.append((num, row))
            except ValueError:
                continue

        for comp in competencies:
            copying = False
            current_elements = []

            for el in document_elements:
                if el.tag.endswith('p'):
                    p = self._element_to_paragraph(el)
                    txt = p.text.strip().replace(" ", "")
                    if txt == comp['code']:
                        copying = True
                        continue
                    elif competency_code_pattern.match(txt) and copying:
                        break
                if copying:
                    current_elements.append(el)

            instruction_pattern = re.compile(r'^(\d+)\.\s*(Инструкция:|Фабула:)')
            instruction_numbers = []
            for el in current_elements:
                if el.tag.endswith('p'):
                    p = self._element_to_paragraph(el)
                    text = p.text.strip()
                    m = instruction_pattern.match(text)
                    if m:
                        instruction_numbers.append(int(m.group(1)))

            num_text = comp['row'].cells[-1].text.strip()
            m = re.match(r'(\d+)-(\d+)', num_text)
            if not m:
                raise Exception(
                    f"Компетенция {comp['code']}: неверный формат диапазона номеров '{num_text}' (должно быть например - 1-16)"
                )
            start = int(m.group(1))
            end = int(m.group(2))
            expected_count = end - start + 1

            if len(instruction_numbers) != expected_count:
                raise Exception(
                    f"Компетенция {comp['code']}: количество заданий ({len(instruction_numbers)}) "
                    f"не совпадает с ожидаемым ({expected_count})"
                )

            new_doc = Document()

            t1 = new_doc.add_table(rows=1, cols=len(first_table.columns))
            for i, c in enumerate(first_table.rows[0].cells):
                t1.rows[0].cells[i].text = c.text
            r = t1.add_row()
            for i, c in enumerate(comp['row'].cells):
                r.cells[i].text = c.text
            self.set_table_borders(t1)

            new_doc.add_paragraph("\n")

            nums = list(range(start, end + 1))
            t2 = new_doc.add_table(rows=1, cols=len(second_table.columns))
            for i, c in enumerate(second_table.rows[0].cells):
                t2.rows[0].cells[i].text = c.text
            for n, r in task_rows:
                if n in nums:
                    row = t2.add_row()
                    for i, c in enumerate(r.cells):
                        row.cells[i].text = c.text
            self.set_table_borders(t2)

            new_doc.add_paragraph("\n")
            heading = new_doc.add_paragraph("Перечень заданий")
            heading.style = 'Heading 2'
            new_doc.add_paragraph("\n")

            for el in current_elements:
                new_doc.element.body.append(deepcopy(el))

            filename = f"{comp['code']}_{original_filename}"
            new_doc.save(os.path.join(output_dir, filename))

    def set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        tblBorders = OxmlElement('w:tblBorders')

        borders = [
            ('top', 'single', 4, '000000'),
            ('left', 'single', 4, '000000'),
            ('bottom', 'single', 4, '000000'),
            ('right', 'single', 4, '000000'),
            ('insideH', 'single', 4, '000000'),
            ('insideV', 'single', 4, '000000')
        ]

        for border_type, border_style, border_size, border_color in borders:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _element_to_paragraph(self, element):
        from docx.oxml import parse_xml
        from docx.text.paragraph import Paragraph
        return Paragraph(parse_xml(element.xml), None)


class SummaryBuilderTab(QWidget):
    def __init__(self):
        super().__init__()
        self.comp_order = {'УК': 1, 'ОПК': 2, 'ПК': 3}
        self.summary_data = []
        self.all_tasks = []
        self.task_mapping = {}
        self.comp_indicators = {}
        self.selected_folder = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        # Форма с полями ввода
        form_layout = QFormLayout()
        form_layout.setFieldGrowthPolicy(QFormLayout.AllNonFixedFieldsGrow)
        form_layout.setLabelAlignment(Qt.AlignLeft)
        form_layout.setFormAlignment(Qt.AlignLeft)
        form_layout.setSpacing(15)

        # Шрифты
        font = QFont()
        font.setFamily("Arial")
        font.setPointSize(12)

        input_font = QFont()
        input_font.setFamily("Arial")
        input_font.setPointSize(14)

        # Стили
        input_style = """
        QLineEdit {
            font-size: 16px;
            padding: 10px;
            min-height: 45px;
            min-width: 400px;
            selection-background-color: #90C8F6;
            border: 1px solid #ccc;
            border-radius: 4px;
        }
        QLineEdit:focus {
            border: 2px solid #5D9CEC;
        }
        """

        button_style = """
        QPushButton {
            font-size: 14px;
            padding: 12px 20px;
            min-height: 45px;
            min-width: 220px;
            background-color: #5D9CEC;
            color: white;
            border: none;
            border-radius: 4px;
        }
        QPushButton:hover {
            background-color: #4A89DC;
        }
        QPushButton:pressed {
            background-color: #3B7DDD;
        }
        QPushButton:disabled {
            background-color: #CCD1D9;
        }
        QPushButton.folder-selected {
            background-color: #A0D468;
        }
        """

        # Поля ввода
        self.direction_input = QLineEdit()
        self.direction_input.setFont(input_font)
        self.direction_input.setStyleSheet(input_style)

        self.profile_input = QLineEdit()
        self.profile_input.setFont(input_font)
        self.profile_input.setStyleSheet(input_style)

        self.year_input = QLineEdit()
        self.year_input.setFont(input_font)
        self.year_input.setStyleSheet(input_style)
        self.year_input.setMaximumWidth(200)

        form_layout.addRow(QLabel("Направление:", font=font), self.direction_input)
        form_layout.addRow(QLabel("Профиль:", font=font), self.profile_input)
        form_layout.addRow(QLabel("Год начала подготовки:", font=font), self.year_input)

        # Кнопка выбора папки и метка пути
        self.select_btn = QPushButton("Выбрать папку с компетенциями")
        self.select_btn.setFont(font)
        self.select_btn.setStyleSheet(button_style)
        self.select_btn.clicked.connect(self.select_directory)

        self.folder_path_label = QLabel("")
        self.folder_path_label.setFont(font)
        self.folder_path_label.setWordWrap(True)
        self.folder_path_label.setStyleSheet("""
            QLabel {
                color: #666;
                font-size: 12px;
                padding: 8px;
                border: 1px solid #eee;
                border-radius: 4px;
                background-color: #f9f9f9;
                min-height: 30px;
            }
        """)

        folder_selection_layout = QVBoxLayout()
        folder_selection_layout.addWidget(self.select_btn)
        folder_selection_layout.addWidget(self.folder_path_label)
        folder_selection_layout.setSpacing(10)

        # Кнопка построения
        self.build_btn = QPushButton("Построить сводный файл")
        self.build_btn.setFont(font)
        self.build_btn.setStyleSheet(button_style)
        self.build_btn.clicked.connect(self.build_summary)
        self.build_btn.setEnabled(False)

        # Прогресс-бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setMinimumHeight(30)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                font-size: 14px;
                min-height: 30px;
                text-align: center;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QProgressBar::chunk {
                background-color: #5D9CEC;
            }
        """)

        # Статусная метка
        self.status_label = QLabel("")
        self.status_label.setFont(font)
        self.status_label.setAlignment(Qt.AlignCenter)

        # Компоновка
        layout.addLayout(form_layout)
        layout.addLayout(folder_selection_layout)
        layout.addWidget(self.build_btn)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)

        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(25)

        self.setLayout(layout)

    def select_directory(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Выберите папку с файлами компетенций")
        if dir_path:
            self.selected_folder = dir_path
            self.select_btn.setProperty("class", "folder-selected")
            self.select_btn.style().polish(self.select_btn)
            self.select_btn.setText("✓ Папка выбрана")
            self.folder_path_label.setText(f"Выбрано: {dir_path}")
            self.build_btn.setEnabled(True)
            self.process_directory(dir_path)
            QMessageBox.information(self, "Успех",
                                  f"Обработано {len(self.summary_data)} дисциплин!\n"
                                  "Теперь можно построить сводный файл.")

    def process_directory(self, dir_path):
        self.summary_data = []
        self.all_tasks = []
        self.comp_indicators = {}

        docx_files = [f for f in os.listdir(dir_path) if f.endswith('.docx')]

        self.progress_bar.setVisible(True)
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(0)
        self.progress_bar.setFormat("Чтение файлов...")
        self.status_label.setText("")
        QApplication.processEvents()

        for i, filename in enumerate(docx_files, start=1):
            self.status_label.setText(f"Обрабатывается файл {i} из {len(docx_files)}: {filename}")
            QApplication.processEvents()
            file_path = os.path.join(dir_path, filename)
            self.process_competency_file(file_path)

        self.progress_bar.setVisible(False)
        self.status_label.setText("")

    def process_competency_file(self, file_path):
        doc = Document(file_path)
        comp_code = os.path.basename(file_path).split('_')[0]

        if len(doc.tables) >= 2:
            first_table = doc.tables[0]

            indicators_text = ""
            header_cells = first_table.rows[0].cells
            indicator_col_idx = None
            for idx, cell in enumerate(header_cells):
                if "Наименование индикаторов" in cell.text:
                    indicator_col_idx = idx
                    break

            if indicator_col_idx is not None:
                indicator_parts = []
                for row in first_table.rows[1:]:
                    text = row.cells[indicator_col_idx].text.strip()
                    if text and text not in indicator_parts:
                        indicator_parts.append(text)
                indicators_text = "\n".join(indicator_parts).strip()

            if indicators_text:
                self.comp_indicators[comp_code] = indicators_text

            for row_idx, row in enumerate(first_table.rows):
                if row_idx == 0:
                    continue
                cells = row.cells
                if len(cells) < 6:
                    continue
                discipline = cells[3].text.strip()
                semester = cells[4].text.strip()
                tasks = cells[5].text.strip()

                self.summary_data.append({
                    'comp_code': comp_code,
                    'discipline': discipline,
                    'semester': semester,
                    'tasks': tasks,
                    'file_path': file_path
                })

            second_table = doc.tables[1]
            for row_idx, row in enumerate(second_table.rows):
                if row_idx == 0:
                    continue
                cells = row.cells
                if len(cells) < 6:
                    continue
                if re.match(r'^\d+\.', cells[0].text.strip()):
                    self.all_tasks.append({
                        'file_path': file_path,
                        'original_num': cells[0].text.strip().split('.')[0],
                        'text': cells[0].text.strip(),
                        'cells': [cell.text.strip() for cell in cells]
                    })

        tasks_section = []
        found_section = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "Перечень заданий" in text:
                found_section = True
                continue
            if found_section and text:
                tasks_section.append(text)

        if tasks_section:
            task_text = "\n".join(tasks_section)
            self.all_tasks.append({
                'file_path': file_path,
                'text': task_text,
                'is_text_section': True
            })

    def build_summary(self):
        if not self.summary_data or not self.all_tasks:
            QMessageBox.warning(self, "Ошибка", "Нет данных для построения сводного файла")
            return

        sorted_data = sorted(
            self.summary_data,
            key=lambda x: (
                self.get_comp_order(x['comp_code']),
                self.parse_semester(x['semester']),
                x['discipline']
            )
        )

        summary_doc = Document()
        self.add_template_header(summary_doc)
        self.add_first_table(summary_doc, sorted_data)
        self.add_second_table(summary_doc, sorted_data)
        self.add_tasks_list(summary_doc, sorted_data)

        save_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить сводный файл", "", "Word Files (*.docx)"
        )
        if save_path:
            summary_doc.save(save_path)
            QMessageBox.information(self, "Готово",
                                    f"Сводный файл успешно создан:\n{save_path}")

    def add_template_header(self, doc):
        def add_centered_bold_paragraph(text):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1  # межстрочный интервал
            p.paragraph_format.space_before = Pt(0)  # интервал перед
            p.paragraph_format.space_after = Pt(0)  # интервал после

            run = p.add_run(text)
            run.bold = True
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            return p

        add_centered_bold_paragraph('Фонд оценочных средств')

        direction = self.direction_input.text().strip() or "____________________"
        profile = self.profile_input.text().strip() or "____________________"
        year = self.year_input.text().strip() or "20__"

        add_centered_bold_paragraph('для оценки остаточных знаний обучающихся по направлению подготовки')
        add_centered_bold_paragraph(f'Направление: {direction}')
        add_centered_bold_paragraph(f'Профиль: {profile}')
        add_centered_bold_paragraph(f'Год начала подготовки – {year}')
        doc.add_paragraph()

    def add_first_table(self, doc, sorted_data):
        doc.add_heading('Распределение тестовых заданий по компетенциям и дисциплинам', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки таблицы (выравнивание по центру)
        headers = table.rows[0].cells
        headers[0].text = "Код компетенции"
        headers[1].text = "Наименование компетенции"
        headers[2].text = "Наименование индикаторов"
        headers[3].text = "Наименование дисциплины/модуля/практики"
        headers[4].text = "Семестр"
        headers[5].text = "Номер задания"

        for cell in headers:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.bold = True

        current_task_num = 1
        comp_groups = defaultdict(list)
        for disc in sorted_data:
            comp_groups[disc['comp_code']].append(disc)

        row_idx = 1
        for comp_code, disciplines in comp_groups.items():
            start_row = row_idx
            for disc in disciplines:
                row_cells = table.add_row().cells
                row_cells[0].text = comp_code if row_idx == start_row else ""
                row_cells[1].text = ""
                row_cells[2].text = (
                    self.comp_indicators.get(comp_code, "")
                    if row_idx == start_row else ""
                )
                row_cells[3].text = disc['discipline']
                row_cells[4].text = disc['semester']

                task_count = self.calculate_task_count(disc['tasks'])
                row_cells[5].text = f"{current_task_num}-{current_task_num + task_count - 1}"

                # Выравнивание всего текста в ячейках по левому краю
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                self.task_mapping[disc['file_path']] = {
                    'discipline': disc['discipline'],
                    'start': current_task_num,
                    'end': current_task_num + task_count - 1,
                    'tasks': []
                }

                current_task_num += task_count
                row_idx += 1

            if len(disciplines) > 1:
                for col in [0, 1, 2]:
                    cell_to_merge = table.cell(start_row, col)
                    for r in range(start_row + 1, row_idx):
                        cell_to_merge.merge(table.cell(r, col))
                    for paragraph in cell_to_merge.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_second_table(self, doc, sorted_data):
        doc.add_heading('Распределение заданий по типам и уровням сложности', level=2)
        doc.add_heading('Ключи к оцениванию', level=3)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки таблицы (выравнивание по центру)
        headers = table.rows[0].cells
        headers[0].text = "№ задания"
        headers[1].text = "Верный ответ"
        headers[2].text = "Критерии"
        headers[3].text = "Тип задания"
        headers[4].text = "Уровень сложности"
        headers[5].text = "Время выполнения (мин.)"

        for cell in headers:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.bold = True

        tasks_by_file = defaultdict(list)
        for task in self.all_tasks:
            if not task.get('is_text_section'):
                tasks_by_file[task['file_path']].append(task)

        current_task_num = 1

        for disc in sorted_data:
            # Заголовок дисциплины (выравнивание по центру)
            row = table.add_row().cells
            row[0].merge(row[5])
            row[0].text = disc['discipline']
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.bold = True

            file_tasks = tasks_by_file.get(disc['file_path'], [])
            task_count_in_first_table = (
                    self.task_mapping[disc['file_path']]['end'] -
                    self.task_mapping[disc['file_path']]['start'] + 1
            )

            for idx in range(task_count_in_first_table):
                row_cells = table.add_row().cells
                new_num = current_task_num + idx
                row_cells[0].text = str(new_num)

                if idx < len(file_tasks):
                    task = file_tasks[idx]
                    for i in range(1, min(6, len(task['cells']))):
                        row_cells[i].text = task['cells'][i]
                else:
                    for i in range(1, 6):
                        row_cells[i].text = "—"

                # Выравнивание всего текста в ячейках по левому краю
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            current_task_num += task_count_in_first_table

    def add_tasks_list(self, doc, sorted_data):
        doc.add_heading('Перечень заданий', level=2)

        for disc in sorted_data:
            source_doc = Document(disc['file_path'])
            found_section = False
            current_num = self.task_mapping[disc['file_path']]['start']

            doc.add_heading(disc['discipline'], level=3)

            for element in source_doc.element.body:
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, source_doc)
                    text = paragraph.text.strip()

                    if "Перечень заданий" in text:
                        found_section = True
                        continue

                    if found_section and text:
                        match = re.match(r'^(\d+)\.\s*(Инструкция:|Фабула:)', text)
                        if match:
                            old_num = match.group(1)
                            new_text = re.sub(r'^(\d+)\.', f'{current_num}.', text, count=1)
                            new_paragraph = doc.add_paragraph()
                            new_paragraph.add_run(new_text).bold = True
                            current_num += 1
                        else:
                            new_paragraph = doc.add_paragraph()
                            for run in paragraph.runs:
                                new_run = new_paragraph.add_run(run.text)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                if run.font.size:
                                    new_run.font.size = run.font.size

                elif element.tag.endswith('tbl'):
                    if found_section:
                        table = Table(element, source_doc)
                        new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = 'Table Grid'

                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_cell = new_table.cell(i, j)
                                new_cell.text = ''
                                for paragraph in cell.paragraphs:
                                    new_paragraph = new_cell.add_paragraph()
                                    for run in paragraph.runs:
                                        new_run = new_paragraph.add_run(run.text)
                                        new_run.bold = run.bold
                                        new_run.italic = run.italic
                                        new_run.underline = run.underline
                                        if run.font.size:
                                            new_run.font.size = run.font.size

                elif element.tag.endswith('drawing'):
                    if found_section:
                        for rel in source_doc.part.rels.values():
                            if "image" in rel.target_ref:
                                image_part = rel.target_part
                                image_bytes = image_part.blob
                                doc.add_picture(io.BytesIO(image_bytes))
                                doc.add_paragraph()

    def merge_cells(self, table, start_row, end_row, col_idx):
        cell_start = table.cell(start_row, col_idx)
        for row in range(start_row + 1, end_row + 1):
            cell_next = table.cell(row, col_idx)
            cell_start.merge(cell_next)

    def get_comp_order(self, code):
        match = re.match(r'([А-Я]+)-(\d+)', code)
        if match:
            comp_type = match.group(1)
            comp_num = int(match.group(2))
            return (self.comp_order.get(comp_type, 99), comp_num)
        return (99, 99)

    def parse_semester(self, semester_str):
        try:
            clean_str = re.sub(r'[^\d,-]', '', semester_str)

            # Если есть запятые (несколько семестров)
            if ',' in clean_str:
                semesters = [int(s.strip()) for s in clean_str.split(',') if s.strip()]
                return min(semesters) + 0.5  # Добавляем 0.5, чтобы диапазон был после одиночного семестра

            # Если есть дефис (диапазон)
            elif '-' in clean_str:
                parts = clean_str.split('-')
                if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                    return int(parts[0]) + 0.5  # Добавляем 0.5, чтобы диапазон был после одиночного семестра

            # Одиночный семестр
            elif clean_str.isdigit():
                return int(clean_str)

            return 0
        except ValueError:
            return 0

    def calculate_task_count(self, tasks_str):
        clean_str = re.sub(r'[^\d,-]', '', tasks_str)
        if '-' in clean_str:
            parts = clean_str.split('-')
            if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                return int(parts[1]) - int(parts[0]) + 1
        elif ',' in clean_str:
            return len(clean_str.split(','))
        try:
            return int(clean_str)
        except ValueError:
            return 0


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ФОС: Разделение и сборка компетенций")
        self.setGeometry(300, 300, 650, 300)

        layout = QVBoxLayout()

        tabs = QTabWidget()
        self.splitter_tab = CompetencySplitterTab()
        self.builder_tab = SummaryBuilderTab()

        tabs.addTab(self.splitter_tab, "Разделение ФОС")
        tabs.addTab(self.builder_tab, "Сборка сводного ФОС")

        layout.addWidget(tabs)
        self.setLayout(layout)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())