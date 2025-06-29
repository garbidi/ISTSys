import sys
import re
import os
import io
from collections import defaultdict
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Length
from docx.shared import Mm
from docx.shared import Emu
from docx.shared import Cm
from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QFileDialog,
    QVBoxLayout, QMessageBox, QHBoxLayout, QProgressBar, QLabel, QLineEdit
)



class SummaryBuilder(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.comp_order = {'УК': 1, 'ОПК': 2, 'ПК': 3}
        self.summary_data = []
        self.all_tasks = []
        self.task_mapping = {}
        self.comp_indicators = {}

    def init_ui(self):
        self.setWindowTitle("Построение сводного файла компетенций")
        self.setGeometry(300, 300, 600, 300)  # чуть увеличим по высоте

        layout = QVBoxLayout()
        input_layout = QHBoxLayout()

        # Метки и поля ввода
        self.direction_label = QLabel("Направление:")
        self.direction_input = QLineEdit()

        self.profile_label = QLabel("Профиль:")
        self.profile_input = QLineEdit()

        self.year_label = QLabel("Год начала подготовки:")
        self.year_input = QLineEdit()

        input_layout.addWidget(self.direction_label)
        input_layout.addWidget(self.direction_input)
        input_layout.addWidget(self.profile_label)
        input_layout.addWidget(self.profile_input)
        input_layout.addWidget(self.year_label)
        input_layout.addWidget(self.year_input)

        btn_layout = QHBoxLayout()
        self.select_btn = QPushButton("Выбрать папку с компетенциями", self)
        self.select_btn.clicked.connect(self.select_directory)

        self.build_btn = QPushButton("Построить сводный файл", self)
        self.build_btn.clicked.connect(self.build_summary)
        self.build_btn.setEnabled(False)

        btn_layout.addWidget(self.select_btn)
        btn_layout.addWidget(self.build_btn)

        self.progress_bar = QProgressBar(self)
        self.progress_bar.setVisible(False)

        self.status_label = QLabel("", self)

        layout.addLayout(input_layout)
        layout.addLayout(btn_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)

        self.setLayout(layout)

    def select_directory(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Выберите папку с файлами компетенций")
        if dir_path:
            self.process_directory(dir_path)
            self.build_btn.setEnabled(True)
            QMessageBox.information(self, "Успех",
                                    f"Обработано {len(self.summary_data)} дисциплин!\n"
                                    "Теперь можно построить сводный файл.")

    def process_directory(self, dir_path):
        self.summary_data = []
        self.all_tasks = []
        self.comp_indicators = {}

        docx_files = [f for f in os.listdir(dir_path) if f.endswith('.docx')]

        self.progress_bar.setVisible(True)
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(0)  # "неопределенный" режим
        self.progress_bar.setFormat("Чтение файлов...")
        self.status_label.setText("")
        QApplication.processEvents()

        for i, filename in enumerate(docx_files, start=1):
            self.status_label.setText(f"Обрабатывается файл {i} из {len(docx_files)}: {filename}")
            QApplication.processEvents()
            file_path = os.path.join(dir_path, filename)
            self.process_competency_file(file_path)

        self.progress_bar.setVisible(False)
        self.status_label.setText("")

    def process_competency_file(self, file_path):
        doc = Document(file_path)
        comp_code = os.path.basename(file_path).split('_')[0]

        if len(doc.tables) >= 2:
            first_table = doc.tables[0]

            indicators_text = ""
            header_cells = first_table.rows[0].cells
            indicator_col_idx = None
            for idx, cell in enumerate(header_cells):
                if "Наименование индикаторов" in cell.text:
                    indicator_col_idx = idx
                    break

            if indicator_col_idx is not None:
                indicator_parts = []
                for row in first_table.rows[1:]:
                    text = row.cells[indicator_col_idx].text.strip()
                    if text and text not in indicator_parts:
                        indicator_parts.append(text)
                indicators_text = "\n".join(indicator_parts).strip()

            if indicators_text:
                self.comp_indicators[comp_code] = indicators_text

            for row_idx, row in enumerate(first_table.rows):
                if row_idx == 0:
                    continue
                cells = row.cells
                if len(cells) < 6:
                    continue
                discipline = cells[3].text.strip()
                semester = cells[4].text.strip()
                tasks = cells[5].text.strip()

                self.summary_data.append({
                    'comp_code': comp_code,
                    'discipline': discipline,
                    'semester': semester,
                    'tasks': tasks,
                    'file_path': file_path
                })

            second_table = doc.tables[1]
            for row_idx, row in enumerate(second_table.rows):
                if row_idx == 0:
                    continue
                cells = row.cells
                if len(cells) < 6:
                    continue
                if re.match(r'^\d+\.', cells[0].text.strip()):
                    self.all_tasks.append({
                        'file_path': file_path,
                        'original_num': cells[0].text.strip().split('.')[0],
                        'text': cells[0].text.strip(),
                        'cells': [cell.text.strip() for cell in cells]
                    })

        tasks_section = []
        found_section = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "Перечень заданий" in text:
                found_section = True
                continue
            if found_section and text:
                tasks_section.append(text)

        if tasks_section:
            task_text = "\n".join(tasks_section)
            self.all_tasks.append({
                'file_path': file_path,
                'text': task_text,
                'is_text_section': True
            })

    def build_summary(self):
        if not self.summary_data or not self.all_tasks:
            QMessageBox.warning(self, "Ошибка", "Нет данных для построения сводного файла")
            return

        sorted_data = sorted(
            self.summary_data,
            key=lambda x: (
                self.get_comp_order(x['comp_code']),
                self.parse_semester(x['semester']),
                x['discipline']
            )
        )

        summary_doc = Document()
        self.add_template_header(summary_doc)
        self.add_first_table(summary_doc, sorted_data)
        self.add_second_table(summary_doc, sorted_data)
        self.add_tasks_list(summary_doc, sorted_data)

        save_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить сводный файл", "", "Word Files (*.docx)"
        )
        if save_path:
            summary_doc.save(save_path)
            QMessageBox.information(self, "Готово",
                                    f"Сводный файл успешно создан:\n{save_path}")

    def add_template_header(self, doc):
        def add_centered_bold_paragraph(text):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = 1  # межстрочный интервал
            p.paragraph_format.space_before = Pt(0)  # интервал перед
            p.paragraph_format.space_after = Pt(0)  # интервал после

            run = p.add_run(text)
            run.bold = True
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            return p

        # Заголовок
        add_centered_bold_paragraph('Фонд оценочных средств')

        # Далее три параграфа с вводом пользователя
        direction = self.direction_input.text().strip() or "____________________"
        profile = self.profile_input.text().strip() or "____________________"
        year = self.year_input.text().strip() or "20__"

        add_centered_bold_paragraph('для оценки остаточных знаний обучающихся по направлению подготовки')
        add_centered_bold_paragraph(f'Направление: {direction}')
        add_centered_bold_paragraph(f'Профиль: {profile}')
        add_centered_bold_paragraph(f'Год начала подготовки – {year}')
        doc.add_paragraph()

    def add_first_table(self, doc, sorted_data):
        doc.add_heading('Распределение тестовых заданий по компетенциям и дисциплинам', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = table.rows[0].cells
        headers[0].text = "Код компетенции"
        headers[1].text = "Наименование компетенции"
        headers[2].text = "Наименование индикаторов"
        headers[3].text = "Наименование дисциплины/модуля/практики"
        headers[4].text = "Семестр"
        headers[5].text = "Номер задания"

        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_task_num = 1
        comp_groups = defaultdict(list)
        for disc in sorted_data:
            comp_groups[disc['comp_code']].append(disc)

        row_idx = 1
        for comp_code, disciplines in comp_groups.items():
            start_row = row_idx
            for disc in disciplines:
                row_cells = table.add_row().cells
                row_cells[0].text = comp_code if row_idx == start_row else ""
                row_cells[1].text = ""
                row_cells[2].text = (
                    self.comp_indicators.get(comp_code, "")
                    if row_idx == start_row else ""
                )
                row_cells[3].text = disc['discipline']
                row_cells[4].text = disc['semester']

                task_count = self.calculate_task_count(disc['tasks'])
                row_cells[5].text = f"{current_task_num}-{current_task_num + task_count - 1}"

                self.task_mapping[disc['file_path']] = {
                    'discipline': disc['discipline'],
                    'start': current_task_num,
                    'end': current_task_num + task_count - 1,
                    'tasks': []
                }

                current_task_num += task_count
                row_idx += 1

            if len(disciplines) > 1:
                for col in [0, 1, 2]:
                    cell_to_merge = table.cell(start_row, col)
                    for r in range(start_row + 1, row_idx):
                        cell_to_merge.merge(table.cell(r, col))
                    for paragraph in cell_to_merge.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_second_table(self, doc, sorted_data):
        doc.add_heading('Распределение заданий по типам и уровням сложности', level=2)
        doc.add_heading('Ключи к оцениванию', level=3)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = table.rows[0].cells
        headers[0].text = "№ задания"
        headers[1].text = "Верный ответ"
        headers[2].text = "Критерии"
        headers[3].text = "Тип задания"
        headers[4].text = "Уровень сложности"
        headers[5].text = "Время выполнения (мин.)"

        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Группируем задания по файлам
        tasks_by_file = defaultdict(list)
        for task in self.all_tasks:
            if not task.get('is_text_section'):
                tasks_by_file[task['file_path']].append(task)

        current_task_num = 1

        for disc in sorted_data:
            file_tasks = tasks_by_file.get(disc['file_path'], [])

            # Добавляем заголовок дисциплины (объединённая строка)
            row = table.add_row().cells
            row[0].merge(row[5])
            row[0].text = disc['discipline']
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Определяем, сколько заданий должно быть (из первой таблицы)
            task_count_in_first_table = (
                    self.task_mapping[disc['file_path']]['end'] -
                    self.task_mapping[disc['file_path']]['start'] + 1
            )

            # Добавляем задания (либо реальные, либо пустые)
            for idx in range(task_count_in_first_table):
                row_cells = table.add_row().cells

                # Номер задания (сквозная нумерация)
                new_num = current_task_num + idx
                row_cells[0].text = str(new_num)

                # Если есть описание задания — заполняем данные
                if idx < len(file_tasks):
                    task = file_tasks[idx]
                    for i in range(1, min(6, len(task['cells']))):
                        row_cells[i].text = task['cells'][i]
                else:
                    # Если описания нет — ставим прочерки
                    for i in range(1, 6):
                        row_cells[i].text = "—"

            # Обновляем текущий номер для следующей дисциплины
            current_task_num += task_count_in_first_table

    def add_tasks_list(self, doc, sorted_data):
        doc.add_heading('Перечень заданий', level=2)

        for disc in sorted_data:
            source_doc = Document(disc['file_path'])
            found_section = False
            current_num = self.task_mapping[disc['file_path']]['start']

            doc.add_heading(disc['discipline'], level=3)

            for element in source_doc.element.body:
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, source_doc)
                    text = paragraph.text.strip()

                    if "Перечень заданий" in text:
                        found_section = True
                        continue

                    if found_section and text:
                        match = re.match(r'^(\d+)\.\s*Инструкция:', text)
                        if match:
                            old_num = match.group(1)
                            new_text = re.sub(r'^(\d+)\.', f'{current_num}.', text, count=1)
                            new_paragraph = doc.add_paragraph()
                            new_paragraph.add_run(new_text).bold = True
                            current_num += 1
                        else:
                            new_paragraph = doc.add_paragraph()
                            for run in paragraph.runs:
                                new_run = new_paragraph.add_run(run.text)
                                new_run.bold = run.bold
                                new_run.italic = run.italic
                                new_run.underline = run.underline
                                if run.font.size:
                                    new_run.font.size = run.font.size

                elif element.tag.endswith('tbl'):
                    if found_section:
                        table = Table(element, source_doc)
                        # Создаем таблицу с правильным количеством строк
                        new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = 'Table Grid'

                        # Копируем содержимое ячеек
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_cell = new_table.cell(i, j)
                                # Очищаем стандартный параграф в новой ячейке
                                new_cell.text = ''
                                # Копируем все параграфы из исходной ячейки
                                for paragraph in cell.paragraphs:
                                    new_paragraph = new_cell.add_paragraph()
                                    for run in paragraph.runs:
                                        new_run = new_paragraph.add_run(run.text)
                                        new_run.bold = run.bold
                                        new_run.italic = run.italic
                                        new_run.underline = run.underline
                                        if run.font.size:
                                            new_run.font.size = run.font.size

                elif element.tag.endswith('drawing'):
                    if found_section:
                        for rel in source_doc.part.rels.values():
                            if "image" in rel.target_ref:
                                image_part = rel.target_part
                                image_bytes = image_part.blob
                                doc.add_picture(io.BytesIO(image_bytes))
                                doc.add_paragraph()

    def merge_cells(self, table, start_row, end_row, col_idx):
        cell_start = table.cell(start_row, col_idx)
        for row in range(start_row + 1, end_row + 1):
            cell_next = table.cell(row, col_idx)
            cell_start.merge(cell_next)

    def get_comp_order(self, code):
        match = re.match(r'([А-Я]+)-(\d+)', code)
        if match:
            comp_type = match.group(1)
            comp_num = int(match.group(2))
            return (self.comp_order.get(comp_type, 99), comp_num)
        return (99, 99)

    def parse_semester(self, semester_str):
        try:
            clean_str = re.sub(r'[^\d-]', '', semester_str)
            if '-' in clean_str:
                parts = clean_str.split('-')
                return int(parts[0])
            return int(clean_str) if clean_str else 0
        except ValueError:
            return 0

    def calculate_task_count(self, tasks_str):
        clean_str = re.sub(r'[^\d,-]', '', tasks_str)
        if '-' in clean_str:
            parts = clean_str.split('-')
            if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                return int(parts[1]) - int(parts[0]) + 1
        elif ',' in clean_str:
            return len(clean_str.split(','))
        try:
            return int(clean_str)
        except ValueError:
            return 0

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = SummaryBuilder()
    window.show()
    sys.exit(app.exec_())